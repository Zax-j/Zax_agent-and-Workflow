import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
import os
import numpy as np
from PIL import Image
import easyocr
import pandas as pd
from io import BytesIO
import json
import re
import base64
from duckduckgo_search import DDGS

from neo4j import GraphDatabase
import docx
import PyPDF2
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# ------------------- 全局配置 -------------------
st.set_page_config(
    page_title="Zax的智能体",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ------------------- 自定义CSS -------------------
st.markdown("""
<style>
.stApp {
    background-color: #171923;
    color: #ffffff;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
}
</style>
""", unsafe_allow_html=True)

# ------------------- 初始化 -------------------
load_dotenv()

# 模型切换核心函数
def get_current_client_and_model():
    model_choice = st.session_state.get("model_selector", "DeepSeek-V4")
    if model_choice == "DeepSeek-V4":
        return OpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            base_url=os.getenv("OPENAI_BASE_URL")
        ), os.getenv("MODEL_NAME")
    elif model_choice == "通义千问VL-Flash(图像理解)":
        return OpenAI(
            api_key=os.getenv("QWEN_VL_API_KEY"),
            base_url=os.getenv("QWEN_VL_BASE_URL")
        ), os.getenv("QWEN_VL_MODEL")

# 初始化OCR
@st.cache_resource
def init_ocr():
    return easyocr.Reader(['ch_sim', 'en'])
ocr = init_ocr()

# 初始化Neo4j
@st.cache_resource
def get_neo4j_driver():
    try:
        return GraphDatabase.driver("bolt://localhost:7687", auth=("neo4j", "123456789"))
    except:
        return None
driver = get_neo4j_driver()

# 联网搜索
def web_search(query):
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=5))
        return "\n".join([f"🔍 {r['body']}" for r in results])
    except:
        return "联网搜索失败"

# 水浒人物列表
SHUIHU_NAMES = [
    "宋江", "卢俊义", "吴用", "公孙胜", "关胜", "林冲", "秦明", "呼延灼", "花荣", "柴进",
    "李应", "朱仝", "鲁智深", "武松", "董平", "张清", "杨志", "徐宁", "索超", "戴宗",
    "刘唐", "李逵", "史进", "穆弘", "雷横", "李俊", "阮小二", "张横", "阮小五", "张顺",
    "阮小七", "杨雄", "石秀", "解珍", "解宝", "燕青", "朱武", "黄信", "孙立", "宣赞",
    "郝思文", "韩滔", "彭玘", "单廷圭", "魏定国", "萧让", "裴宣", "欧鹏", "邓飞", "燕顺",
    "杨林", "凌振", "蒋敬", "吕方", "郭盛", "安道全", "皇甫端", "王英", "扈三娘", "鲍旭",
    "樊瑞", "孔明", "孔亮", "项充", "李衮", "金大坚", "马麟", "童威", "童猛", "孟康",
    "侯健", "陈达", "杨春", "郑天寿", "陶宗旺", "宋清", "乐和", "龚旺", "丁得孙", "穆春",
    "曹正", "薛永", "施恩", "李忠", "周通", "汤隆", "杜兴", "邹渊", "邹润", "朱贵",
    "朱富", "蔡福", "蔡庆", "李立", "李云", "焦挺", "石勇", "孙新", "顾大嫂", "张青",
    "孙二娘", "王定六", "郁保四", "白胜", "时迁", "段景住"
]

# Neo4j查询
def query_neo4j(question):
    mentioned = [n for n in SHUIHU_NAMES if n in question]
    if len(mentioned) < 2 or not driver:
        return ""
    res = []
    with driver.session() as session:
        for name in mentioned:
            recs = session.run("MATCH (a:Hero{name:$name})-[r]-(b:Hero) RETURN a.name, type(r), b.name", name=name)
            for r in recs:
                res.append(f"{r['a.name']} 与 {r['b.name']} 的关系：{r['type(r)']}")
    return "\n".join(res)

# Excel导出
def generate_excel_download(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return st.download_button("📥 一键下载Excel", output, "表格.xlsx")

# 文档读取
def read_pdf(file):
    return "\n".join(p.extract_text() for p in PyPDF2.PdfReader(file).pages)
def read_word(file):
    return "\n".join(p.text for p in docx.Document(file).paragraphs)

# RAG构建
def build_rag(text):
    docs = [Document(page_content=text)]
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return FAISS.from_documents(splitter.split_documents(docs), embeddings)
def retrieve_context(q, db):
    if not db:
        return ""
    return "\n".join([d.page_content for d in db.similarity_search(q, k=3)])

# 图片转Base64
def encode_image_to_base64(image_file):
    img_bytes = image_file.getvalue()
    base64_str = base64.b64encode(img_bytes).decode("utf-8")
    return f"data:image/jpeg;base64,{base64_str}"

# 导出对话历史
def export_chat_history():
    md_content = "# Zax智能体对话记录\n\n"
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            md_content += f"**用户**: {msg['content']}\n\n"
        elif msg["role"] == "assistant":
            md_content += f"**助手**: {msg['content']}\n\n"
        md_content += "---\n"
    return md_content

# ------------------- 会话状态初始化 -------------------
if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "system", "content": "你是Zax的智能助手，可根据用户开启的工具回答问题，简洁专业。"}]

for k in ["ocr_text", "edited_ocr_text", "vector_db", "last_uploaded_doc", "last_uploaded_img"]:
    if k not in st.session_state:
        st.session_state[k] = None
for k in ["vl_image", "last_uploaded_vl_image"]:
    if k not in st.session_state:
        st.session_state[k] = None

# 工具开关
if "enable_rag" not in st.session_state: st.session_state.enable_rag = True
if "enable_neo4j" not in st.session_state: st.session_state.enable_neo4j = True
if "enable_ocr" not in st.session_state: st.session_state.enable_ocr = True
if "enable_web" not in st.session_state: st.session_state.enable_web = False
if "enable_vl" not in st.session_state: st.session_state.enable_vl = False

if "model_selector" not in st.session_state:
    st.session_state.model_selector = "DeepSeek-V4"

# ------------------- 聊天界面（直接默认进入） -------------------
with st.sidebar:
    st.title("⚙️ Harness Engineering")
    st.divider()

    # 模型选择 + 场景说明
    st.subheader("🤖 大模型切换")
    st.selectbox(
        "选择当前使用模型",
        ["DeepSeek-V4", "通义千问VL-Flash(图像理解)"],
        key="model_selector"
    )
    if st.session_state.model_selector == "DeepSeek-V4":
        st.success("💡 通用模型：文本对话、文档问答、逻辑推理、表格生成")
    else:
        st.success("💡 视觉模型：仅用于图片内容理解")
    st.divider()

    # 文档知识库
    st.subheader("📄 文档知识库")
    up_file = st.file_uploader("上传PDF/Word", type=["pdf", "docx"])
    if up_file:
        last_doc = st.session_state.get("last_uploaded_doc")
        if (last_doc is None or last_doc.name != up_file.name or last_doc.size != up_file.size):
            with st.spinner("🔨 知识库构建中..."):
                txt = read_pdf(up_file) if up_file.name.endswith(".pdf") else read_word(up_file)
                st.session_state.vector_db = build_rag(txt)
                st.session_state.last_uploaded_doc = up_file
            st.success("✅ 知识库构建完成")
        else:
            st.info(f"📁 已加载：{up_file.name}")
    else:
        st.session_state.vector_db = None
        st.session_state.last_uploaded_doc = None

    st.divider()

    # OCR图片 + 编辑框
    st.subheader("🖼️ OCR图片")
    up_img = st.file_uploader("上传图片（表格生成）", type=["jpg", "png", "jpeg"], key="ocr_upload")
    if up_img:
        last_img = st.session_state.get("last_uploaded_img")
        if (last_img is None or last_img.name != up_img.name or last_img.size != up_img.size):
            with st.spinner("🔍 OCR识别中..."):
                img = Image.open(up_img)
                st.session_state.ocr_text = " ".join(ocr.readtext(np.array(img), detail=0))
                st.session_state.edited_ocr_text = st.session_state.ocr_text
                st.session_state.last_uploaded_img = up_img
            st.success("✅ OCR识别完成，可在下方编辑文本")
        else:
            st.info(f"🖼️ 已加载：{up_img.name}")

    # OCR文本编辑框（直接放在上传下方）
    if st.session_state.ocr_text:
        st.subheader("✏️ OCR文本编辑")
        st.session_state.edited_ocr_text = st.text_area(
            "识别结果（可直接修改）",
            value=st.session_state.edited_ocr_text if st.session_state.edited_ocr_text else st.session_state.ocr_text,
            height=150
        )
        st.info("💡 修改后，发送「把OCR生成的结果制作成表格并导出」即可使用编辑后的内容")
    else:
        st.info("上传图片后，识别结果将显示在这里供你编辑")

    st.divider()

    # 图像理解
    st.subheader("🖼️ 图像理解")
    vl_image = st.file_uploader("上传图片（理解内容）", type=["jpg", "png", "jpeg"], key="vl_upload")
    if vl_image:
        st.session_state.vl_image = vl_image
        st.success("✅ 图片已加载")
    else:
        st.session_state.vl_image = None

    st.divider()

    # 工具开关
    st.subheader("🛠️ 工具开关")
    st.session_state.enable_rag = st.toggle("文档知识库", value=st.session_state.enable_rag)
    st.session_state.enable_neo4j = st.toggle("水浒知识图谱", value=st.session_state.enable_neo4j)
    st.session_state.enable_ocr = st.toggle("OCR表格生成", value=st.session_state.enable_ocr)
    st.session_state.enable_web = st.toggle("联网搜索", value=st.session_state.enable_web)
    st.session_state.enable_vl = st.toggle("图像理解", value=st.session_state.enable_vl)

st.title("Zax的智能体 - 对话界面")
st.caption("支持：文档问答 | OCR表格生成 | 水浒人物关系 | 联网搜索 | 图像理解")

# 渲染历史消息
for msg in st.session_state.messages:
    if msg["role"] != "system":
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

prompt = st.chat_input("输入你的问题...")
if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        tool_logs = []
        context = ""
        client, current_model = get_current_client_and_model()

        # 工具调用逻辑
        if st.session_state.enable_rag and st.session_state.vector_db:
            doc_ctx = retrieve_context(prompt, st.session_state.vector_db)
            context += f"【文档内容】\n{doc_ctx}\n\n"
            tool_logs.append("✅ 调用文档知识库")
        if st.session_state.enable_neo4j:
            neo_ctx = query_neo4j(prompt)
            if neo_ctx:
                context += f"【人物关系】\n{neo_ctx}\n\n"
                tool_logs.append("✅ 调用水浒知识图谱")
        if st.session_state.enable_web:
            web_ctx = web_search(prompt)
            context += f"【联网搜索】\n{web_ctx}\n\n"
            tool_logs.append("✅ 调用联网搜索")
        # 使用用户编辑后的OCR文本
        if st.session_state.enable_ocr and st.session_state.edited_ocr_text and any(k in prompt for k in ["表格", "excel", "导出"]):
            context += f"【OCR文本】\n{st.session_state.edited_ocr_text}\n\n请整理为JSON格式（仅包含headers和rows字段）"
            tool_logs.append("✅ 调用OCR表格生成（使用编辑后的文本）")
        if st.session_state.enable_vl and st.session_state.vl_image:
            context += f"【图片内容】已上传图片，用户需求：{prompt}"
            tool_logs.append("✅ 调用图像理解")

        # 可视化工具日志
        with st.expander("🔧 工具调用日志", expanded=True):
            for log in tool_logs:
                st.write(log)

        # 消息构造
        if st.session_state.enable_vl and st.session_state.vl_image:
            img_b64 = encode_image_to_base64(st.session_state.vl_image)
            messages = [
                {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": img_b64}},
                    {"type": "text", "text": prompt}
                ]}
            ]
        else:
            messages = st.session_state.messages + [
                {"role": "user", "content": f"参考资料：\n{context}\n用户问题：{prompt}"}
            ]

        # 第一步：先显示spinner，调用API获取stream（只包裹请求过程）
        with st.spinner("思考中..."):
            stream = client.chat.completions.create(
                model=current_model,
                messages=messages,
                stream=True
            )

        # 第二步：spinner消失后，再开始流式输出（不再和spinner重叠）
        reply = st.write_stream(stream)

        # OCR表格解析
        df = None
        show_reply = True
        if st.session_state.enable_ocr and st.session_state.edited_ocr_text and any(k in prompt for k in ["表格", "excel", "导出"]):
            match = re.search(r'\{[^{}]*"headers"\s*:\s*\[[^\]]*\][^{}]*"rows"\s*:\s*\[[^\]]*\][^{}]*\}', str(reply), re.DOTALL)
            if match:
                try:
                    data = json.loads(match.group())
                    df = pd.DataFrame(data["rows"], columns=data["headers"])
                    st.dataframe(df, use_container_width=True)
                    generate_excel_download(df)
                    show_reply = False
                except:
                    pass

    st.session_state.messages.append({"role": "assistant", "content": str(reply)})

# 导出对话按钮（移到最后，确保包含最新消息）
st.divider()
col1, col2 = st.columns([1, 4])
with col1:
    chat_md = export_chat_history()
    st.download_button("📥 导出对话", chat_md, "对话记录.md", use_container_width=True)


# ====================== API 调用专用核心函数（无界面，纯逻辑）======================
def zax_agent_core(
    prompt: str,
    # 工具开关
    enable_rag: bool = True,
    enable_neo4j: bool = True,
    enable_ocr: bool = True,
    enable_web: bool = False,
    enable_vl: bool = False,
    # 外部传入数据（替代streamlit session_state）
    vector_db = None,
    edited_ocr_text: str = None,
    vl_image = None,
    # 模型选择
    model_choice: str = "DeepSeek-V4"
):
    """
    Zax智能体核心逻辑（API专用，无界面依赖）
    :return: 最终回答文本
    """
    try:
        tool_logs = []
        context = ""

        # 1. 获取模型客户端
        if model_choice == "DeepSeek-V4":
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"), base_url=os.getenv("OPENAI_BASE_URL"))
            current_model = os.getenv("MODEL_NAME")
        elif model_choice == "通义千问VL-Flash(图像理解)":
            client = OpenAI(api_key=os.getenv("QWEN_VL_API_KEY"), base_url=os.getenv("QWEN_VL_BASE_URL"))
            current_model = os.getenv("QWEN_VL_MODEL")
        else:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"), base_url=os.getenv("OPENAI_BASE_URL"))
            current_model = os.getenv("MODEL_NAME")

        # 2. 拼接工具上下文
        if enable_rag and vector_db:
            doc_ctx = retrieve_context(prompt, vector_db)
            context += f"【文档内容】\n{doc_ctx}\n\n"
            tool_logs.append("调用文档知识库")
        if enable_neo4j:
            neo_ctx = query_neo4j(prompt)
            if neo_ctx:
                context += f"【人物关系】\n{neo_ctx}\n\n"
                tool_logs.append("调用水浒知识图谱")
        if enable_web:
            web_ctx = web_search(prompt)
            context += f"【联网搜索】\n{web_ctx}\n\n"
            tool_logs.append("调用联网搜索")
        if enable_ocr and edited_ocr_text and any(k in prompt for k in ["表格", "excel", "导出"]):
            context += f"【OCR文本】\n{edited_ocr_text}\n\n请整理为JSON格式（仅包含headers和rows字段）"
            tool_logs.append("调用OCR表格生成")

        # 3. 构造消息
        if enable_vl and vl_image:
            img_b64 = encode_image_to_base64(vl_image)
            messages = [
                {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": img_b64}},
                    {"type": "text", "text": prompt}
                ]}
            ]
        else:
            messages = [
                {"role": "system", "content": "你是Zax的智能助手，可根据用户开启的工具回答问题，简洁专业。"},
                {"role": "user", "content": f"参考资料：\n{context}\n用户问题：{prompt}"}
            ]

        # 4. 调用大模型（非流式，适配API）
        response = client.chat.completions.create(
            model=current_model,
            messages=messages,
            stream=False
        )
        reply = response.choices[0].message.content

        return reply

    except Exception as e:
        return f"智能体执行异常：{str(e)}"