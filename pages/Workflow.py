import streamlit as st
from dotenv import load_dotenv
import os
from openai import OpenAI
from typing import TypedDict, List
import re
from docx import Document
from io import BytesIO
from docx.oxml.ns import qn  # 关键修复：导入字体设置需要的函数

# ========================= 环境配置 =========================
load_dotenv()
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url=os.getenv("OPENAI_BASE_URL")
)
model_name = os.getenv("MODEL_NAME")

# ========================= 会话状态初始化 =========================
if "step" not in st.session_state:
    st.session_state.step = 0  # 0:初始 1:清洗完成 2:润色完成 3:报告完成
if "raw_text" not in st.session_state:
    st.session_state.raw_text = ""
if "cleaned_text" not in st.session_state:
    st.session_state.cleaned_text = ""
if "polished_text" not in st.session_state:
    st.session_state.polished_text = ""
if "final_report" not in st.session_state:
    st.session_state.final_report = ""
# 新增：记录上次上传的文件，避免重复处理
if "last_uploaded_file" not in st.session_state:
    st.session_state.last_uploaded_file = None


# ========================= 核心功能函数 =========================
# 1. 双层日志清洗（规则+LLM）
def clean_logs(raw: str) -> str:
    # 规则清洗：去乱码、符号、空行、重复
    raw = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', raw)
    raw = raw.replace("，", "\n").replace("。", "\n").replace("；", "\n")
    raw = re.sub(r'\s+', ' ', raw)
    raw = re.sub(r'[!！?？~#$%^&*]+', '', raw)
    lines = [line.strip() for line in raw.split("\n") if line.strip() and len(line) > 1]
    lines = list(dict.fromkeys(lines))  # 去重

    if not lines:
        return ""

    # LLM清洗：纠错、标准化
    prompt = f"""
    清洗工作记录：修正错别字、口语化、乱码，保留原意，每条一行，不增删内容。
    原始内容：
    {chr(10).join(lines)}
    """
    res = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}],
                                         temperature=0.1)
    return res.choices[0].message.content.strip()


# 2. 内容润色（无强制分类，连贯书面化）
def polish_text(text: str) -> str:
    prompt = f"""
    将以下工作记录润色为正式、连贯的书面化工作内容，不拆分分类，保持段落通顺：
    {text}
    """
    res = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}],
                                         temperature=0.2)
    return res.choices[0].message.content.strip()


# 3. 差异化模板生成（核心优化！）
def generate_report(text: str, template: str) -> str:
    if template == "标准周报":
        prompt = f"""
        根据工作内容生成【标准周报】，结构：
        一、本周工作完成情况
        二、问题与解决方案
        三、下周工作计划
        内容正式完整，使用中文标点，不要使用列表符号
        工作内容：{text}
        """
    elif template == "研发工程师版":
        prompt = f"""
        根据工作内容生成【研发周报】，结构：
        一、开发进度与技术实现
        二、问题排查与风险处理
        三、优化方案与后续安排
        侧重技术细节，专业严谨，使用中文标点，不要使用列表符号
        工作内容：{text}
        """
    else:  # 极简日报
        prompt = f"""
        根据工作内容生成【极简日报】，结构：
        今日工作内容
        明日工作计划
        简洁干练，不冗余，使用中文标点，不要使用列表符号
        工作内容：{text}
        """

    res = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}],
                                         temperature=0.3)
    return res.choices[0].message.content.strip()


# 4. 修复版：Markdown转Word
def md_to_docx(md_text: str):
    # 先清理文本中的特殊字符和控制字符
    cleaned_text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', md_text)
    cleaned_text = re.sub(r'\u200B|\u200C|\u200D|\uFEFF', '', cleaned_text)  # 去除零宽空格等

    doc = Document()
    # 设置中文字体，避免乱码
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lines = cleaned_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            heading = doc.add_heading(line[2:], 1)
            heading.style.font.name = '宋体'
            heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif line.startswith("## "):
            heading = doc.add_heading(line[3:], 2)
            heading.style.font.name = '宋体'
            heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif line.startswith("### "):
            heading = doc.add_heading(line[4:], 3)
            heading.style.font.name = '宋体'
            heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif line.startswith("- "):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            para.style.font.name = '宋体'
            para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            para = doc.add_paragraph(line)
            para.style.font.name = '宋体'
            para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


# ========================= Streamlit 界面 =========================
st.set_page_config(page_title="智能周报生成器", layout="wide")
st.title("📋 智能工作日志生成器")
st.caption("✅ 分步编辑 | ✅ 差异化模板 | ✅ 文件上传 | ✅ Word导出（无乱码）")

# ---------- 输入区域：文本 + TXT上传（修复循环刷新） ----------
st.subheader("1. 输入工作日志")
upload_file = st.file_uploader("上传 .txt 日志文件", type="txt")

# 输入框和模板选择
input_col1, input_col2 = st.columns([3, 1])
with input_col1:
    # 文本框直接绑定session_state，修改后自动保存
    raw_text = st.text_area("日志内容", value=st.session_state.raw_text, height=180)
with input_col2:
    template = st.selectbox("选择报告模板", ["标准周报", "研发工程师版", "极简日报"])

# 文件上传逻辑：只在文件变化时更新，不触发rerun
if upload_file is not None:
    # 检查文件是否和上次上传的相同，不同才更新
    if (st.session_state.last_uploaded_file is None or
            upload_file.name != st.session_state.last_uploaded_file.name or
            upload_file.size != st.session_state.last_uploaded_file.size):
        # 读取文件内容并更新session_state
        st.session_state.raw_text = upload_file.read().decode("utf-8")
        # 记录当前文件信息，避免重复处理
        st.session_state.last_uploaded_file = upload_file
        # 关键修复：去掉st.rerun()，文本框会自动在下一次渲染时更新

st.divider()

# ---------- 分步操作区域（核心：单步推进 + 单框编辑） ----------
st.subheader("2. 分步处理流程")
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("🔍 第一步：清洗日志", disabled=st.session_state.step >= 1, use_container_width=True):
        with st.spinner("双层清洗中..."):
            st.session_state.cleaned_text = clean_logs(raw_text)
            st.session_state.step = 1
with col2:
    if st.button("✍️ 第二步：润色内容", disabled=st.session_state.step < 1, use_container_width=True):
        with st.spinner("内容润色中..."):
            st.session_state.polished_text = polish_text(st.session_state.cleaned_text)
            st.session_state.step = 2
with col3:
    if st.button("📄 第三步：生成报告", disabled=st.session_state.step < 2, use_container_width=True):
        with st.spinner("生成报告中..."):
            st.session_state.final_report = generate_report(st.session_state.polished_text, template)
            st.session_state.step = 3

# ---------- 一键生成模式 ----------
if st.button("🚀 一键生成完整报告", type="primary", use_container_width=True):
    with st.spinner("正在全自动处理..."):
        st.session_state.cleaned_text = clean_logs(raw_text)
        st.session_state.polished_text = polish_text(st.session_state.cleaned_text)
        st.session_state.final_report = generate_report(st.session_state.polished_text, template)
        st.session_state.step = 3

st.divider()

# ---------- 分步结果展示（单框编辑，简洁美观） ----------
st.subheader("3. 处理结果")
if st.session_state.step >= 1:
    st.markdown("**✅ 清洗完成（可直接修改）**")
    st.session_state.cleaned_text = st.text_area("清洗后内容", st.session_state.cleaned_text, height=150,
                                                 label_visibility="collapsed")

if st.session_state.step >= 2:
    st.markdown("**✅ 润色完成（可直接修改）**")
    st.session_state.polished_text = st.text_area("润色后内容", st.session_state.polished_text, height=150,
                                                  label_visibility="collapsed")

if st.session_state.step >= 3:
    st.markdown("**✅ 报告生成完成（可直接修改）**")
    st.session_state.final_report = st.text_area("最终报告", st.session_state.final_report, height=300,
                                                 label_visibility="collapsed")

    # 双格式导出（修复Word乱码）
    export_col1, export_col2 = st.columns(2)
    with export_col1:
        st.download_button("💾 导出 Markdown", st.session_state.final_report, "报告.md", use_container_width=True)
    with export_col2:
        try:
            doc = md_to_docx(st.session_state.final_report)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("💾 导出 Word", buf, "报告.docx", use_container_width=True)
        except Exception as e:
            st.error(f"Word导出失败: {str(e)}")

# ---------- 重置按钮 ----------
if st.button("🔄 重置所有内容", use_container_width=True):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()